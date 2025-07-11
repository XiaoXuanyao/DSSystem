
import json
import torch
import chromadb
import random
import time
import docx
from debug import *
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
from langchain_core.documents import Document
from openai import OpenAI
from transformers import AutoTokenizer, AutoModelForCausalLM, AutoModelForSeq2SeqLM, pipelines



"""
LLM API Key
"""
api_key = json.load(open("key.json"))["api_key"]



def load_pdf(
        path: str
    ) -> PyPDFLoader:
    """
    从path加载PDF文档

    Args:
        path (str): PDF文件的路径

    Returns:
        PyPDFLoader: 加载的PDF文档对象

    """
    pdf_abs_path = os.path.abspath(path)

    if not os.path.exists(pdf_abs_path):
        raise FileNotFoundError(f"未找到 PDF 文件：{pdf_abs_path}")

    # 加载 PDF 文档
    loader = PyPDFLoader(pdf_abs_path)
    return loader



def split_pdf(
        loader: PyPDFLoader,
        filename: str,
        chunk_size: int=512,
        chunk_overlap: int=128
    ) -> list[Document]:
    """
    将PDF文档切分为更小的块

    Args:
        pdf_loader (PyPDFLoader): 加载的PDF文档对象
        filename (str): PDF文件的名称

    Returns:
        list: 切分后的文档块列表
    
    """
    documents = loader.load()
    debug.log(f"Lib", f"从 PDF 加载了 {len(documents)} 页。")
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len
    )
    # 将文档切分为更小的块
    chunks = text_splitter.split_documents(documents)
    debug.log(f"Lib", f"PDF 已切分为 {len(chunks)} 个块。")
    for i, chunk in enumerate(chunks):
        chunk.metadata["source"] = filename
        chunk.metadata["chunk_index"] = i + 1
    return chunks



def load_docx(
        path: str
    ) -> docx.Document:
    """
    从path加载DOCX文档

    Args:
        path (str): DOCX文件的路径

    Returns:
        Document: 加载的DOCX文档对象

    """
    docx_abs_path = os.path.abspath(path)

    if not os.path.exists(docx_abs_path):
        raise FileNotFoundError(f"未找到 DOCX 文件：{docx_abs_path}")

    # 加载 DOCX 文档
    doc = docx.Document(docx_abs_path)
    return doc



def split_docx(
        doc: docx.Document,
        filename:str,
        chunk_size: int=512,
        chunk_overlap: int=128
    ) -> list[Document]:
    """
    将DOCX文档切分为更小的块

    Args:
        doc (docx.Document): 加载的DOCX文档对象
    
    Returns:
        list: 切分后的文档块列表
    
    """
    content = []
    for para in doc.paragraphs:
        content.append(para.text)
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        length_function=len
    )
    chunks = text_splitter.split_text("\n".join(content))
    debug.log(f"Lib", f"DOCX 已切分为 {len(chunks)} 个块。")
    for i in range(len(chunks)):
        chunk = chunks[i]
        doc = Document(
            page_content=chunk,
            metadata={
                "source": filename,
                "chunk_index": i + 1,
            }
        )
        chunks[i] = doc
    return chunks



def load_embedding_model(
        model_name: str = None
    ) -> SentenceTransformer:
    """
    加载嵌入模型

    Args:
        model_name (str): 嵌入模型的名称

    Returns:
        SentenceTransformer: 加载的嵌入模型对象

    """
    if model_name is None:
        model_name = "BAAI/bge-small-zh-v1.5"
    if torch.cuda.is_available():
        device = "cuda"
        debug.log(f"Lib", f"CUDA is available! Using GPU for computation.")
    else:
        device = "cpu"
        debug.log(f"Lib", f"CUDA is NOT available. Using CPU for computation.")

    model = SentenceTransformer(
        model_name,
        device = device,
        local_files_only = True
    )
    debug.log(f"Lib", f"成功加载模型: {model_name} 到 {device}。")
    return model



def load_summary_model(
        model_name: str="fnlp/bart-base-chinese"
    ) -> tuple[AutoTokenizer, AutoModelForSeq2SeqLM]:
    """
    加载摘要模型

    Args:
        model_name (str): 摘要模型的名称，默认为 "fnlp/bart-base-chinese"

    Returns:
        AutoModelForSeq2SeqLM: 加载的摘要模型对象
    """
    tokenizer = AutoTokenizer.from_pretrained(model_name, local_files_only=True)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name, torch_dtype=torch.bfloat16, local_files_only=True)
    return tokenizer, model



def load_local_llm(
        model_path: str
    ) -> pipelines.pipeline:
    """
    加载本地LLM模型

    Args:
        model_path (str): LLM模型的路径

    Returns:
        LocalLLM: 加载的本地LLM模型对象
    """
    tokenizer = AutoTokenizer.from_pretrained(model_path)
    model = AutoModelForCausalLM.from_pretrained(model_path, torch_dtype=torch.bfloat16)
    llm = pipelines.pipeline(
        "text-generation",
        model=model,
        tokenizer=tokenizer,
        max_length=512,
        temperature=0.1,
        top_p=0.9,
        pad_token_id=tokenizer.eos_token_id,
        eos_token_id=tokenizer.eos_token_id,
    )
    return llm



def load_deepseek_online() -> OpenAI:
    """
    返回 DeepSeek V1 API
    """
    return OpenAI(
        api_key=api_key,
        base_url="https://api.deepseek.com/v1"
    )



def store_chunks(
        model: SentenceTransformer,
        chunks: list[Document],
        path: str,
        collection_name: str="sentence_embeddings"
    ):
    """
    存储句子嵌入至向量数据库

    Args:
        model (SentenceTransformer): 嵌入模型对象
        chunks (list[Document]): 需要嵌入的文档块列表
        path (str): 存储路径
        collection_name (str): 向量数据库集合名称，默认为 "sentence_embeddings"

    """
    sentences = [chunk.page_content for chunk in chunks]
    embeddings = model.encode(sentences, convert_to_tensor=True)
    debug.log(f"Lib", f"句子嵌入的形状 (句子数, 维度): {embeddings.shape}")
    debug.log(f"Lib", f"嵌入张量所在的设备: {embeddings.device}")  # 打印张量当前所在的设备

    Client = chromadb.PersistentClient(path=path)  # 使用持久化客户端
    existing_collections = [col.name for col in Client.list_collections()]
    if collection_name in existing_collections:
        debug.log(f"Lib", f"集合 '{collection_name}' 已存在，正在删除旧集合...")
        Client.delete_collection(name=collection_name)
    collection = Client.create_collection(name=collection_name)
    
    embeddings = embeddings.cpu().numpy().tolist()  # 转换为列表格式
    collection.add(
        ids=[f"{random.randint(100_000_000, 999_999_999):06d}-{random.randint(100_000_000, 999_999_999):06d}" for _ in embeddings],
        embeddings=embeddings,
        documents=sentences,
        metadatas=[chunk.metadata for chunk in chunks]
    )



def gen_summary(
        text: str,
        tokenizer: AutoTokenizer,
        model: AutoModelForSeq2SeqLM,
        max_length: int=-1
    ) -> str:
    """
    生成文本摘要

    Args:
        model (AutoModelForSeq2SeqLM): 摘要模型对象
        tokenizer (AutoTokenizer): 摘要模型的分词器
        text (str): 输入文本
        max_length (int): 摘要的最大长度，默认为 100

    Returns:
        str: 生成的摘要文本

    """
    if max_length == -1:
        max_length = len(text) // 2 + 25
    inputs = tokenizer(
            text,
            truncation=True,
            max_length=512,
            padding="max_length",
            return_tensors="pt"
        )
    refsum = model.generate(
        inputs["input_ids"],
        max_length=max_length,
        min_length=5,
        num_beams=2,
    )
    refsum = tokenizer.decode(refsum[0], skip_special_tokens=True)
    return refsum.replace(" ", "")



def query_message(
        model: SentenceTransformer,
        path: str,
        collection_name: str,
        query: str,
        n_results: int=2
    ) -> dict:
    """
    查询向量数据库中的嵌入

    Args:
        path (str): 向量数据库的存储路径
        collection_name (str): 向量数据库集合名称
        query (str): 查询字符串
        n_results (int): 返回的最相似结果数量，默认为 2

    Returns:
        list: 查询结果列表

    """
    Client = chromadb.PersistentClient(path=path)
    collection = Client.get_collection(name=collection_name)
    query_embedding = model.encode([query], convert_to_tensor=True).cpu().numpy().tolist()
    results = collection.query(
        query_embeddings=query_embedding,
        n_results=n_results  # 返回最相似的 n_results 个结果
    )
    return results



def document_to_string(
        doc_results: dict,
        use_summary: bool=True,
        summary_tokenizer: AutoTokenizer=None,
        summary_model: AutoModelForSeq2SeqLM=None
    ) -> str:
    """
    将 Document 对象转换为字符串

    Args:
        doc (Document): 要转换的 Document 对象

    Returns:
        str: 文档内容字符串

    """
    res = ""
    for i in range(len(doc_results['documents'][0])):
        if doc_results['distances'][0][i] > 0.8:
            continue

        text = doc_results['documents'][0][i]
        if use_summary and summary_tokenizer is not None and summary_model is not None:
            text = gen_summary(text, summary_tokenizer, summary_model)
        metadata = str(doc_results['metadatas'][0][i])
        if use_summary and summary_tokenizer is not None and summary_model is not None:
            metadata = gen_summary(metadata, summary_tokenizer, summary_model)

        res += "{\n"
        res += f"text: {text}\n"
        res += f"metadata: {metadata}\n"
        res += f"distance: {doc_results['distances'][0][i]}\n"
        res += "},\n"
    return res